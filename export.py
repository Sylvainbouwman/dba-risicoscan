"""
Word-export voor het DBA Risicoscan-memo.
"""

from __future__ import annotations
from io import BytesIO
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


SIGNAAL_LABELS = {
    "loondienst": "Wijst op loondienst",
    "neutraal": "Neutraal",
    "zzp": "Wijst op ZZP-schap",
}

STERKTE_LABELS = {
    "zwak": "zwak",
    "matig": "matig",
    "sterk": "sterk",
}


def _stel_stijl_in(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _voeg_koptekst_toe(doc: Document, intake: dict) -> None:
    doc.add_heading("DBA Risicoscan – Risicomemo", level=0)

    p = doc.add_paragraph()
    p.add_run(f"Opdrachtgever: ").bold = True
    p.add_run(intake.get("opdrachtgever", "–"))

    p = doc.add_paragraph()
    p.add_run("Opdrachtnemer: ").bold = True
    p.add_run(intake.get("opdrachtnemer", "–"))

    p = doc.add_paragraph()
    p.add_run("Sector: ").bold = True
    p.add_run(intake.get("sector", "–"))

    p = doc.add_paragraph()
    p.add_run("Aard werkzaamheden: ").bold = True
    p.add_run(intake.get("werkzaamheden", "–"))

    p = doc.add_paragraph()
    p.add_run("Startdatum opdracht: ").bold = True
    p.add_run(intake.get("startdatum", "–"))

    p = doc.add_paragraph()
    p.add_run("Datum memo: ").bold = True
    p.add_run(date.today().strftime("%d %B %Y"))

    doc.add_paragraph()


def _voeg_samenvatting_toe(doc: Document, analyse: dict) -> None:
    doc.add_heading("Samenvatting", level=1)
    doc.add_paragraph(analyse.get("patroon_samenvatting", ""))

    risico_signalen = analyse.get("sterke_risicosignalen", [])
    if risico_signalen:
        doc.add_heading("Sterke risicosignalen", level=2)
        for signaal in risico_signalen:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(signaal)

    doc.add_paragraph()


def _voeg_gezichtspunten_toe(doc: Document, analyse: dict) -> None:
    doc.add_heading("Analyse per gezichtspunt", level=1)

    for gp in analyse.get("gezichtspunten_analyse", []):
        signaal = gp.get("signaal", "neutraal")
        sterkte = gp.get("sterkte", "")
        label = SIGNAAL_LABELS.get(signaal, signaal)
        sterkte_label = STERKTE_LABELS.get(sterkte, sterkte)

        heading = doc.add_heading(
            f"{gp.get('nummer')}. {gp.get('naam', '')}", level=2
        )

        p = doc.add_paragraph()
        run = p.add_run(f"Signaal: {label} ({sterkte_label})")
        run.bold = True
        if signaal == "loondienst":
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif signaal == "zzp":
            run.font.color.rgb = RGBColor(0x37, 0x86, 0x10)
        else:
            run.font.color.rgb = RGBColor(0xBF, 0x85, 0x00)

        doc.add_paragraph(gp.get("redenering", ""))

        p = doc.add_paragraph()
        p.add_run("Bron: ").bold = True
        p.add_run(gp.get("bron", ""))

        doc.add_paragraph()


def _voeg_verbeterpunten_toe(doc: Document, analyse: dict) -> None:
    verbeterpunten = analyse.get("verbeterpunten", [])
    if not verbeterpunten:
        return

    doc.add_heading("Verbeterpunten", level=1)

    for vp in verbeterpunten:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{vp.get('gezichtspunt', '')}: ").bold = True
        p.add_run(vp.get("actie", ""))

        if vp.get("toelichting"):
            toel = doc.add_paragraph(vp["toelichting"])
            toel.paragraph_format.left_indent = Pt(18)

    doc.add_paragraph()


def _voeg_disclaimer_toe(doc: Document) -> None:
    doc.add_heading("Disclaimer", level=1)
    disclaimer = (
        "Dit memo is een indicatieve risicoanalyse op basis van de door u verstrekte informatie "
        "en de negen gezichtspunten uit het Deliveroo-arrest (HR 24 maart 2023) en het Uber-arrest "
        "(HR 21 februari 2025). Het memo vormt geen fiscaal of juridisch advies en kan niet worden "
        "gebruikt als zekerheid over de kwalificatie van de arbeidsrelatie. De uiteindelijke "
        "beoordeling is voorbehouden aan de Belastingdienst en de rechter, op basis van alle feiten "
        "en omstandigheden en de feitelijke uitvoering van de opdracht. Raadpleeg bij twijfel een "
        "fiscalist of arbeidsrechtadvocaat."
    )
    p = doc.add_paragraph(disclaimer)
    p.runs[0].italic = True


def genereer_memo(intake: dict, analyse: dict) -> BytesIO:
    """Genereer een Word-document als BytesIO-object."""
    doc = Document()
    _stel_stijl_in(doc)
    _voeg_koptekst_toe(doc, intake)
    _voeg_samenvatting_toe(doc, analyse)
    _voeg_gezichtspunten_toe(doc, analyse)
    _voeg_verbeterpunten_toe(doc, analyse)
    _voeg_disclaimer_toe(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
